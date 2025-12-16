from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document


def read_paragraphs(path: str | Path) -> list[str]:
    import hashlib
    import logging
    
    logger = logging.getLogger(__name__)
    path_obj = Path(path)
    
    logger.info(f"[read_paragraphs] Reading DOCX: {path_obj}")
    logger.info(f"[read_paragraphs] Absolute path: {path_obj.resolve()}")
    
    # Verify file hash before reading
    with open(path_obj, 'rb') as f:
        file_content = f.read()
        file_hash = hashlib.md5(file_content).hexdigest()
    logger.info(f"[read_paragraphs] File hash: {file_hash}")
    logger.info(f"[read_paragraphs] File size: {len(file_content):,} bytes")
    
    document = Document(path_obj)
    paragraphs = [para.text for para in document.paragraphs]
    
    logger.info(f"[read_paragraphs] Extracted {len(paragraphs)} paragraphs")
    if paragraphs:
        logger.info(f"[read_paragraphs] First paragraph: {paragraphs[0][:100]}...")
    
    return paragraphs


def write_paragraphs(source_path: str | Path, translations: Iterable[str], output_path: str | Path) -> None:
    document = Document(source_path)
    for para, translated in zip(document.paragraphs, translations, strict=False):
        para.text = translated
    document.save(output_path)


def write_new_document(translations: Iterable[str], output_path: str | Path) -> None:
    doc = Document()
    for paragraph in translations:
        doc.add_paragraph(paragraph)
    doc.save(output_path)

