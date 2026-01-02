from __future__ import annotations

import hashlib
import json
import logging
import os
import tempfile
import threading
import time
import uuid
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel

from .claude_client import ClaudeTranslator
from .export import export_tbx, export_tmx
from .processing import PDF_SUFFIX, translate_file_to_memory
from .settings import get_settings
from .terminology import Glossary, TranslationMemory

# Configure logging with both console and file handlers
def setup_logging():
    """Configure logging to write to both console and file."""
    log_dir = Path(__file__).parent.parent.parent / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)
    log_file = log_dir / "backend.log"
    
    # Create formatter
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
    
    # Get root logger
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.INFO)
    
    # Remove existing handlers to avoid duplicates
    root_logger.handlers.clear()
    
    # Console handler (stdout)
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(formatter)
    root_logger.addHandler(console_handler)
    
    # File handler (append mode)
    file_handler = logging.FileHandler(log_file, mode='a', encoding='utf-8')
    file_handler.setLevel(logging.INFO)
    file_handler.setFormatter(formatter)
    root_logger.addHandler(file_handler)
    
    return log_file

# Setup logging on module import
log_file_path = setup_logging()
logger = logging.getLogger(__name__)
logger.info(f"Logging initialized. Log file: {log_file_path}")

app = FastAPI(title="Legal Translator API", version="0.1.0")

# CORS configuration - allow all origins in production, specific origins in development
cors_origins = os.getenv("CORS_ORIGINS", "").split(",") if os.getenv("CORS_ORIGINS") else []
if not cors_origins:
    # Default: allow localhost and ngrok for development
    cors_origins = [
        "http://localhost:3000",
        "http://localhost:5173",
    ]
    allow_origin_regex = r"https://.*\.ngrok\.(io|app|free\.app)"
else:
    # In production, use specific origins or allow all if "*" is set
    allow_origin_regex = None
    if "*" in cors_origins:
        cors_origins = ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_origin_regex=allow_origin_regex,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Job storage - each job is completely isolated
translation_jobs: dict[str, dict] = {}

# Request tracking to detect duplicates
request_log: list[dict] = []

JOB_NOT_FOUND = "Job not found"


class TranslationStatus(BaseModel):
    job_id: str
    status: str
    progress: Optional[float] = None
    current_paragraph: Optional[int] = None
    total_paragraphs: Optional[int] = None
    error: Optional[str] = None
    pdf_size: Optional[int] = None  # Size of PDF in bytes (in memory)
    original_text: Optional[list[str]] = None  # Original paragraphs
    translated_text: Optional[list[str]] = None  # Translated paragraphs
    report: Optional[dict] = None


def _load_glossary(glossary_path: Path | None, source_lang: str, target_lang: str) -> Glossary | None:
    if glossary_path and glossary_path.exists():
        return Glossary.from_csv(glossary_path, source_lang=source_lang, target_lang=target_lang, name=glossary_path.stem)
    return None


def find_glossary_files() -> list[Path]:
    glossaries = []
    for dir_path in [Path("glossary"), Path("tests/docs")]:
        if dir_path.exists():
            glossaries.extend(dir_path.glob("*.csv"))
    return sorted(set(glossaries))


def detect_language_from_file(file_path: Path) -> str:
    """Detect language from a document file."""
    try:
        from langdetect import detect_langs, LangDetectException
    except ImportError:
        logger.warning("langdetect not available, falling back to default language")
        return "fr"
    
    try:
        # Extract text from file
        text = ""
        if file_path.suffix.lower() == ".docx":
            from docx import Document
            doc = Document(file_path)
            text = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
        elif file_path.suffix.lower() == ".pdf":
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(page_text)
            text = "\n\n".join(text_parts)
        elif file_path.suffix.lower() == ".txt":
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        
        if not text or len(text.strip()) < 50:
            logger.warning(f"Not enough text to detect language (got {len(text.strip())} chars)")
            return "fr"
        
        # Use more text for better accuracy (langdetect works better with more text)
        # Use up to 5000 characters for better detection accuracy
        sample_text = text[:5000] if len(text) > 5000 else text
        
        # Get language probabilities
        languages = detect_langs(sample_text)
        
        # Map langdetect codes to our language codes
        lang_map = {
            "fr": "fr",  # French
            "de": "de",  # German
            "it": "it",  # Italian
            "en": "en",  # English
        }
        
        # Get the most probable language that we support
        for lang_obj in languages:
            detected_code = lang_obj.lang
            if detected_code in lang_map:
                detected = lang_map[detected_code]
                logger.info(f"Detected language: {detected_code} (prob: {lang_obj.prob:.2f}) -> {detected}")
                return detected
        
        # If none of the detected languages are in our map, use the most probable one
        # and log a warning
        if languages:
            best_match = languages[0].lang
            logger.warning(f"Detected language '{best_match}' not in supported languages, defaulting to 'fr'")
            logger.info(f"All detected languages: {[(l.lang, l.prob) for l in languages[:5]]}")
        
        return "fr"
        
    except LangDetectException as e:
        logger.warning(f"Language detection failed: {e}, using default 'fr'")
        return "fr"
    except Exception as e:
        logger.error(f"Error detecting language: {e}", exc_info=True)
        return "fr"


def _run_translation(job_id: str) -> None:
    """
    Run translation in a completely isolated context.
    All data comes from job storage, no closure variables.
    """
    try:
        # Get job data from storage
        job = translation_jobs.get(job_id)
        if not job:
            logger.error(f"[TRANSLATION {job_id}] Job not found in storage")
            return
        
        logger.info("=" * 80)
        logger.info(f"[TRANSLATION {job_id}] ===== STARTING TRANSLATION THREAD =====")
        logger.info(f"[TRANSLATION {job_id}] Job data from storage:")
        logger.info(f"[TRANSLATION {job_id}]   - File hash: {job.get('file_hash', 'MISSING')}")
        logger.info(f"[TRANSLATION {job_id}]   - Input path: {job.get('input_path', 'MISSING')}")
        logger.info(f"[TRANSLATION {job_id}]   - Source: {job.get('source_lang')} -> Target: {job.get('target_lang')}")
        
        job["status"] = "in_progress"
        
        # Extract all data from job storage
        input_file_path = Path(job["input_path"])
        expected_hash = job["file_hash"]
        source_lang = job["source_lang"]
        target_lang = job["target_lang"]
        glossary_path_str = job.get("glossary_path")
        skip_memory = job.get("skip_memory", False)  # Default False means memory is enabled
        custom_prompt = job.get("custom_prompt")
        user_role = job.get("user_role", "")
        
        # Verify input file exists and hash matches
        if not input_file_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_file_path}")
        
        # Read and verify file content
        logger.info(f"Job {job_id}: [FILE VERIFICATION] Opening file: {input_file_path}")
        logger.info(f"Job {job_id}: [FILE VERIFICATION] File absolute path: {input_file_path.resolve()}")
        logger.info(f"Job {job_id}: [FILE VERIFICATION] File exists: {input_file_path.exists()}")
        
        with open(input_file_path, 'rb') as f:
            file_content = f.read()
            actual_hash = hashlib.md5(file_content).hexdigest()
        
        logger.info(f"Job {job_id}: [FILE VERIFICATION] Read {len(file_content):,} bytes from file")
        logger.info(f"Job {job_id}: [FILE VERIFICATION] Expected hash: {expected_hash}")
        logger.info(f"Job {job_id}: [FILE VERIFICATION] Actual hash: {actual_hash}")
        logger.info(f"Job {job_id}: [FILE VERIFICATION] First 100 bytes: {file_content[:100]}")
        
        if actual_hash != expected_hash:
            logger.error(f"Job {job_id}: [FILE VERIFICATION] HASH MISMATCH!")
            logger.error(f"Job {job_id}: [FILE VERIFICATION] Expected: {expected_hash[:16]}...")
            logger.error(f"Job {job_id}: [FILE VERIFICATION] Got: {actual_hash[:16]}...")
            raise ValueError(f"File hash mismatch! Expected {expected_hash[:16]}..., got {actual_hash[:16]}...")
        
        logger.info(f"Job {job_id}: [FILE VERIFICATION] Hash verified successfully")
        logger.info(f"Job {job_id}: [FILE VERIFICATION] File: {input_file_path.name}, Size: {len(file_content):,} bytes")
        
        # Create completely isolated temporary directory for this job
        job_temp_dir = tempfile.mkdtemp(prefix=f"translation_{job_id}_")
        job["temp_dir"] = job_temp_dir
        logger.info(f"Job {job_id}: Created isolated temp directory: {job_temp_dir}")
        
        try:
            # No output file needed - we'll store PDF in memory
            
            # Load glossary (loaded fresh from file for each translation - no restart needed!)
            glossary = None
            if glossary_path_str:
                glossary_path_obj = Path(glossary_path_str)
                logger.info(f"Job {job_id}: Loading glossary from {glossary_path_obj} (file will be read fresh)")
                glossary = _load_glossary(glossary_path_obj, source_lang, target_lang)
                if glossary:
                    logger.info(f"Job {job_id}: Loaded glossary '{glossary.name}' with {len(glossary)} entries")
                else:
                    logger.warning(f"Job {job_id}: Failed to load glossary from {glossary_path_obj}")
            
            # Use shared translation memory (persisted across jobs)
            settings = get_settings()
            memory_file = settings.data_root / "memory.json"
            memory = TranslationMemory(memory_file)
            logger.info(f"Job {job_id}: Using shared translation memory from {memory_file}")
            logger.info(f"Job {job_id}: Memory contains {len(memory)} existing records")
            
            # Extract translation pairs from reference document if provided
            reference_doc_pairs = {}
            reference_doc_path_str = job.get("reference_doc_path")
            if reference_doc_path_str:
                reference_doc_path_obj = Path(reference_doc_path_str)
                if reference_doc_path_obj.exists():
                    logger.info(f"Job {job_id}: Extracting translation pairs from reference document...")
                    from .processing import extract_translation_pairs_from_reference_doc
                    # Create translator temporarily for extraction
                    temp_translator = ClaudeTranslator(
                        api_key=settings.anthropic_api_key,
                        dry_run=False,
                    )
                    reference_doc_pairs = extract_translation_pairs_from_reference_doc(
                        reference_doc_path=reference_doc_path_obj,
                        translator=temp_translator,
                        source_lang=source_lang,
                        target_lang=target_lang,
                    )
                    logger.info(f"Job {job_id}: Extracted {len(reference_doc_pairs)} translation pairs from reference doc")
                else:
                    logger.warning(f"Job {job_id}: Reference doc path provided but file not found: {reference_doc_path_obj}")
            
            # Determine prompt to use: custom_prompt if provided, otherwise role-specific prompt
            from .claude_client import _load_prompt_template_for_role
            prompt_to_use = custom_prompt
            if not prompt_to_use:
                # Use role-specific prompt if no custom prompt provided
                prompt_to_use = _load_prompt_template_for_role(user_role)
                logger.info(f"Job {job_id}: Using role-specific prompt for role: '{user_role}' (length: {len(user_role) if user_role else 0})")
            else:
                logger.info(f"Job {job_id}: Using custom prompt provided by user")
            
            # Create translator
            translator = ClaudeTranslator(
                api_key=settings.anthropic_api_key,
                dry_run=False,
                custom_prompt_template=prompt_to_use,
            )
            
            # Progress callback
            def update_progress(idx: int, total: int, length: int):
                if job.get("cancelled", False):
                    logger.info(f"Job {job_id}: Cancellation detected")
                    raise KeyboardInterrupt("Translation cancelled")
                
                job["current_paragraph"] = idx
                job["total_paragraphs"] = total
                # Don't set progress to 100% until all paragraphs are done AND post-processing is complete
                # Reserve 10% for post-processing (PDF generation, report building, etc.)
                if idx >= total:
                    # All paragraphs translated, but still processing
                    job["progress"] = 0.90  # 90% - still need to generate PDF and finalize
                else:
                    # Calculate progress as percentage of paragraph translation (90% of total work)
                    job["progress"] = (idx / total) * 0.90 if total > 0 else 0.0
            
            # Run translation - verify file again before calling translate_file_to_memory
            logger.info(f"Job {job_id}: [PRE-TRANSLATE] About to call translate_file_to_memory")
            logger.info(f"Job {job_id}: [PRE-TRANSLATE] Input path: {input_file_path}")
            logger.info(f"Job {job_id}: [PRE-TRANSLATE] Input path absolute: {input_file_path.resolve()}")
            logger.info(f"Job {job_id}: [PRE-TRANSLATE] Input file exists: {input_file_path.exists()}")
            
            # Verify file one more time right before translation
            with open(input_file_path, 'rb') as verify_f:
                verify_content = verify_f.read()
                verify_hash = hashlib.md5(verify_content).hexdigest()
            logger.info(f"Job {job_id}: [PRE-TRANSLATE] Final verification - Hash: {verify_hash}")
            logger.info(f"Job {job_id}: [PRE-TRANSLATE] Expected hash: {expected_hash}")
            if verify_hash != expected_hash:
                logger.error(f"Job {job_id}: [PRE-TRANSLATE] FILE CHANGED BEFORE TRANSLATION!")
                raise ValueError(f"File changed! Expected {expected_hash[:16]}..., got {verify_hash[:16]}...")
            
            logger.info(f"Job {job_id}: [PRE-TRANSLATE] File verified, calling translate_file_to_memory")
            logger.info(f"Job {job_id}: [PRE-TRANSLATE] Source lang: {source_lang}, Target lang: {target_lang}")
            
            # Read original text before translation (using the same function that translate_file_to_memory uses)
            # Import here to avoid circular imports
            import sys
            from pathlib import Path as PathLib
            if input_file_path.suffix.lower() == ".docx":
                from docx import Document
                doc = Document(input_file_path)
                original_text = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
            elif input_file_path.suffix.lower() == ".pdf":
                import pdfplumber
                text_parts = []
                with pdfplumber.open(input_file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            text_parts.append(page_text)
                original_text = "\n\n".join(text_parts)
            elif input_file_path.suffix.lower() == ".txt":
                with open(input_file_path, 'r', encoding='utf-8') as f:
                    original_text = f.read()
            else:
                original_text = ""
            
            # Split original text into paragraphs for display
            original_paragraphs = [p.strip() for p in original_text.split("\n\n") if p.strip()]
            if not original_paragraphs:
                original_paragraphs = [p.strip() for p in original_text.split("\n") if p.strip()]
            if not original_paragraphs:
                original_paragraphs = [original_text] if original_text else []
            
            # Store original text in job for later use
            job["original_text"] = original_paragraphs
            logger.info(f"Job {job_id}: Stored {len(original_paragraphs)} original paragraphs")
            
            pdf_bytes, translated_paragraphs, report = translate_file_to_memory(
                input_path=input_file_path,
                glossary=glossary,
                memory=memory,
                translator=translator,
                source_lang=source_lang,
                target_lang=target_lang,
                progress_callback=update_progress,
                skip_memory=skip_memory,
                reference_doc_pairs=reference_doc_pairs if reference_doc_pairs else None,
            )
            
            logger.info(f"[TRANSLATION {job_id}] Translation complete")
            logger.info(f"[TRANSLATION {job_id}] PDF generated in memory: {len(pdf_bytes):,} bytes")
            logger.info(f"[TRANSLATION {job_id}] Original paragraphs: {len(original_paragraphs)}")
            logger.info(f"[TRANSLATION {job_id}] Translated paragraphs: {len(translated_paragraphs)}")
            
            # Update job status - CRITICAL: Check if job was overwritten
            current_job = translation_jobs.get(job_id)
            if current_job and current_job.get("file_hash") != expected_hash:
                logger.error(f"[TRANSLATION {job_id}] ⚠️  JOB WAS OVERWRITTEN!")
                logger.error(f"[TRANSLATION {job_id}] Expected hash: {expected_hash}")
                logger.error(f"[TRANSLATION {job_id}] Current hash: {current_job.get('file_hash')}")
                raise ValueError("Job was overwritten by another request during translation")
            
            # Store PDF bytes, original text, and translated text in memory (not on filesystem)
            # Original text was already stored in job above
            # Extract applied glossary and memory terms from report
            applied_glossary_terms = report.get("applied_glossary_terms", [])
            applied_memory_terms = report.get("applied_memory_terms", [])
            
            job.update({
                "status": "completed",
                "progress": 1.0,
                "pdf_bytes": pdf_bytes,  # Store PDF in memory
                "pdf_size": len(pdf_bytes),
                "translated_text": translated_paragraphs,  # Store translated paragraphs (with <glossary> and <memory> markers)
                "original_text": job.get("original_text"),  # Store original text for side-by-side comparison
                "applied_glossary_terms": applied_glossary_terms,  # Store glossary terms that were applied
                "applied_memory_terms": applied_memory_terms,  # Store memory terms that were applied
                "report": report,
                "completed_at": time.time(),  # Track when completed
            })
            
            logger.info(f"[TRANSLATION {job_id}] Status updated to completed")
            logger.info(f"[TRANSLATION {job_id}] Final job hash: {job.get('file_hash')}")
            logger.info(f"[TRANSLATION {job_id}] ===== TRANSLATION COMPLETE =====")
            logger.info("=" * 80)
            
        finally:
            # Cleanup: remove input file immediately after translation
            if input_file_path.exists():
                try:
                    input_file_path.unlink()
                    logger.info(f"Job {job_id}: Cleaned up input file")
                except Exception as e:
                    logger.warning(f"Job {job_id}: Failed to cleanup input file: {e}")
            
    except KeyboardInterrupt:
        logger.info(f"Job {job_id}: Translation cancelled")
        job = translation_jobs.get(job_id, {})
        job.update({
            "status": "cancelled",
            "error": "Translation cancelled by user",
        })
    except Exception as e:
        logger.error(f"Job {job_id}: Translation failed: {e}", exc_info=True)
        job = translation_jobs.get(job_id, {})
        job.update({
            "status": "failed",
            "error": str(e),
        })


@app.get("/")
async def root():
    return {"message": "Legal Translator API", "version": "0.1.0"}

@app.get("/api/health")
async def health():
    return {"status": "healthy", "service": "legal-translator-api"}


@app.get("/streaming-demo.html")
async def streaming_demo():
    """Serve the streaming translation demo page."""
    demo_path = Path(__file__).parent.parent.parent / "streaming-demo.html"
    if not demo_path.exists():
        raise HTTPException(status_code=404, detail="Demo page not found")
    return FileResponse(demo_path, media_type="text/html")


@app.get("/api/glossaries")
async def list_glossaries():
    return {"glossaries": [str(g) for g in find_glossary_files()]}


@app.get("/api/glossary/{glossary_name}/content")
async def get_glossary_content(glossary_name: str):
    """Get the content of a glossary file."""
    import csv
    
    # Find the glossary file
    glossary_files = find_glossary_files()
    glossary_path = None
    for g in glossary_files:
        if g.stem == glossary_name or g.name == glossary_name:
            glossary_path = g
            break
    
    if not glossary_path or not glossary_path.exists():
        raise HTTPException(status_code=404, detail=f"Glossary '{glossary_name}' not found")
    
    # Read and return glossary content
    entries = []
    try:
        with glossary_path.open("r", encoding="utf-8-sig", newline="") as handle:
            reader = csv.DictReader(handle)
            if not {"term", "translation"}.issubset(reader.fieldnames or []):
                raise HTTPException(status_code=400, detail="Invalid glossary format")
            for row in reader:
                term = (row.get("term") or "").strip()
                translation = (row.get("translation") or "").strip()
                if term and translation:
                    entries.append({
                        "term": term,
                        "translation": translation,
                        "context": (row.get("context") or "").strip() or None
                    })
    except Exception as e:
        logger.error(f"Error reading glossary {glossary_name}: {e}")
        raise HTTPException(status_code=500, detail=f"Error reading glossary: {str(e)}")
    
    return {
        "name": glossary_path.stem,
        "path": str(glossary_path),
        "entries": entries,
        "total": len(entries)
    }


@app.get("/api/memory/content")
async def get_memory_content():
    """Get all translation memory entries."""
    settings = get_settings()
    memory_file = settings.data_root / "memory.json"

    # Use TranslationMemory to get cleaned entries (filters out stale entries)
    try:
        memory = TranslationMemory(memory_file)
        entries = []
        for record in memory:
            entries.append({
                "key": record.key,
                "source_text": record.source_text,
                "translated_text": record.translated_text,
                "source_lang": record.source_lang,
                "target_lang": record.target_lang
            })
        return {
            "path": str(memory_file),
            "entries": entries,
            "total": len(entries)
        }
    except Exception as e:
        logger.error(f"Error reading memory: {e}")
        raise HTTPException(status_code=500, detail=f"Error reading memory: {str(e)}")


@app.put("/api/memory/content")
async def update_memory_content(request: Request):
    """Update the content of the translation memory file."""
    settings = get_settings()
    memory_file = settings.data_root / "memory.json"
    
    # Ensure parent directory exists
    memory_file.parent.mkdir(parents=True, exist_ok=True)
    
    # Get updated entries from request body
    try:
        body = await request.json()
        entries = body.get("entries", [])
        
        if not isinstance(entries, list):
            raise HTTPException(status_code=400, detail="Invalid entries format")
        
        # Validate entries - skip empty entries (new entries that weren't filled in)
        valid_entries = []
        for entry in entries:
            if not isinstance(entry, dict):
                raise HTTPException(status_code=400, detail="Each entry must be an object")
            source_text = (entry.get("source_text") or "").strip()
            translated_text = (entry.get("translated_text") or "").strip()
            # Skip empty entries (allow empty fields in editing, but filter them out when saving)
            if not source_text or not translated_text:
                continue
            valid_entries.append(entry)
        
        entries = valid_entries
        
        # Convert entries back to memory JSON format (with keys)
        memory_data = {}
        from .terminology.memory import TranslationRecord
        from hashlib import sha1
        
        for entry in entries:
            source_lang = (entry.get("source_lang") or "").strip()
            target_lang = (entry.get("target_lang") or "").strip()
            source_text = (entry.get("source_text") or "").strip()
            translated_text = (entry.get("translated_text") or "").strip()
            
            # Generate key using same method as TranslationRecord
            raw = f"{source_lang}:{target_lang}:{source_text.strip()}"
            key = sha1(raw.encode("utf-8")).hexdigest()
            
            memory_data[key] = {
                "source_text": source_text,
                "translated_text": translated_text,
                "source_lang": source_lang,
                "target_lang": target_lang
            }
        
        # Write updated entries to JSON file
        memory_file.write_text(
            json.dumps(memory_data, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
        
        logger.info(f"Updated memory with {len(entries)} entries")
        
        # Return updated entries in the same format as GET
        return {
            "path": str(memory_file),
            "entries": [
                {
                    "key": key,
                    "source_text": record["source_text"],
                    "translated_text": record["translated_text"],
                    "source_lang": record["source_lang"],
                    "target_lang": record["target_lang"]
                }
                for key, record in memory_data.items()
            ],
            "total": len(memory_data)
        }
        
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON in request body")
    except Exception as e:
        logger.error(f"Error updating memory: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error updating memory: {str(e)}")


@app.delete("/api/memory/content")
async def delete_all_memory_content():
    """Delete all translation memory entries."""
    settings = get_settings()
    memory_file = settings.data_root / "memory.json"
    
    try:
        # Write empty JSON object to clear all entries
        memory_file.parent.mkdir(parents=True, exist_ok=True)
        memory_file.write_text("{}", encoding="utf-8")
        
        logger.info("Deleted all memory entries")
        
        return {
            "path": str(memory_file),
            "entries": [],
            "total": 0,
            "message": "All memory entries deleted successfully"
        }
    except Exception as e:
        logger.error(f"Error deleting memory: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error deleting memory: {str(e)}")


@app.put("/api/glossary/{glossary_name}/content")
async def update_glossary_content(glossary_name: str, request: Request):
    """Update the content of a glossary file."""
    import csv
    
    # Find the glossary file
    glossary_files = find_glossary_files()
    glossary_path = None
    for g in glossary_files:
        if g.stem == glossary_name or g.name == glossary_name:
            glossary_path = g
            break
    
    if not glossary_path or not glossary_path.exists():
        raise HTTPException(status_code=404, detail=f"Glossary '{glossary_name}' not found")
    
    # Get updated entries from request body
    try:
        body = await request.json()
        entries = body.get("entries", [])
        
        if not isinstance(entries, list):
            raise HTTPException(status_code=400, detail="Invalid entries format")
        
        # Validate entries - skip empty entries (new entries that weren't filled in)
        valid_entries = []
        for entry in entries:
            if not isinstance(entry, dict):
                raise HTTPException(status_code=400, detail="Each entry must be an object")
            term = (entry.get("term") or "").strip()
            translation = (entry.get("translation") or "").strip()
            # Skip empty entries (allow empty terms/translations in editing, but filter them out when saving)
            if not term or not translation:
                continue
            valid_entries.append(entry)
        
        entries = valid_entries
        
        # Write updated entries to CSV file
        with glossary_path.open("w", encoding="utf-8-sig", newline="") as handle:
            fieldnames = ["term", "translation", "context"]
            writer = csv.DictWriter(handle, fieldnames=fieldnames)
            writer.writeheader()
            
            for entry in entries:
                writer.writerow({
                    "term": (entry.get("term") or "").strip(),
                    "translation": (entry.get("translation") or "").strip(),
                    "context": (entry.get("context") or "").strip() or ""
                })
        
        logger.info(f"Updated glossary {glossary_name} with {len(entries)} entries")
        
        return {
            "name": glossary_path.stem,
            "path": str(glossary_path),
            "entries": entries,
            "total": len(entries),
            "message": "Glossary updated successfully"
        }
        
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON in request body")
    except Exception as e:
        logger.error(f"Error updating glossary {glossary_name}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error updating glossary: {str(e)}")


def _get_user_role(username: str) -> str:
    """
    Get the user's role from environment variables.
    Checks USER_X_ROLES and USER_X_USERNAME environment variables.
    Returns the first role found, or 'user' as default.
    """
    if not username:
        return "user"
    
    username_lower = username.strip().lower()
    
    # Check USER_X_USERNAME and USER_X_ROLES pairs
    for i in range(1, 100):  # Check up to USER_100
        user_username_key = f"USER_{i}_USERNAME"
        user_roles_key = f"USER_{i}_ROLES"
        
        # Check if username matches
        user_username = os.getenv(user_username_key, "").strip().lower()
        if user_username and user_username == username_lower:
            roles = os.getenv(user_roles_key, "").strip()
            if roles:
                # Return the first role (roles can be comma-separated)
                return roles.split(",")[0].strip()
    
    # Check ADMIN_USERS env var for admin role
    admin_users_env = os.getenv("ADMIN_USERS", "").strip()
    if admin_users_env:
        admin_users = [u.strip().lower() for u in admin_users_env.split(",") if u.strip()]
        if username_lower in admin_users:
            return "admin"
    
    # Fallback: check if username contains "admin"
    if "admin" in username_lower:
        return "admin"
    
    # Default role
    return "user"


def _check_admin_role(request: Request) -> bool:
    """
    Check if the current user has admin role.
    For MVP: checks X-User-Role header or username-based role mapping.
    Returns False by default (deny access) unless explicitly granted admin role.
    """
    # Check for role in header (set by frontend) - this is the primary check
    # FastAPI headers are case-insensitive, try both cases
    user_role = (request.headers.get("X-User-Role") or request.headers.get("x-user-role") or "").strip().lower()
    
    # Explicitly check for "admin" role - must be exact match
    if user_role == "admin":
        logger.debug(f"Admin access granted via X-User-Role header: {user_role}")
        return True
    
    # If role header is present but not "admin", deny access immediately
    if user_role and user_role != "admin":
        logger.debug(f"Access denied - non-admin role specified: '{user_role}'")
        return False
    
    # Fallback: check username from header (if provided and no role header was set)
    # FastAPI headers are case-insensitive, try both cases
    username = (request.headers.get("X-Username") or request.headers.get("x-username") or "").strip()
    if username:
        # Check against ADMIN_USERS env var first
        admin_users_env = os.getenv("ADMIN_USERS", "").strip()
        if admin_users_env:
            admin_users = [u.strip().lower() for u in admin_users_env.split(",") if u.strip()]
            if username.lower() in admin_users:
                logger.debug(f"Admin access granted via ADMIN_USERS env var for user: {username}")
                return True
        
        # Only check username pattern if no explicit admin users list
        # This is a fallback for MVP - username must contain "admin"
        if not admin_users_env and "admin" in username.lower():
            logger.debug(f"Admin access granted via username pattern for user: {username}")
            return True
    
    # No admin role found - deny access (default deny)
    logger.debug(f"Access denied - user_role: '{user_role}', username: '{username}'")
    return False


@app.get("/api/user-role")
async def get_user_role(request: Request):
    """
    Get the user's role based on username.
    Returns the actual role name (e.g., "translator", "reviewer", "admin") instead of just "user" or "admin".
    """
    username = (request.headers.get("X-Username") or request.headers.get("x-username") or "").strip()
    if not username:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Username is required"
        )
    
    role = _get_user_role(username)
    return {"username": username, "role": role}


@app.get("/api/prompt")
async def get_prompt(request: Request):
    """
    Get the translation prompt template for the current user's role.
    Only accessible to users with admin role for viewing.
    Non-admin users have the prompt loaded in background but cannot view it.
    """
    # Log received headers for debugging
    user_role_header = request.headers.get("x-user-role", "NOT_SET")
    username_header = request.headers.get("x-username", "NOT_SET")
    logger.info(f"[PROMPT ACCESS] X-User-Role: {user_role_header}, X-Username: {username_header}")
    
    # Check if user has admin role
    is_admin = _check_admin_role(request)
    logger.info(f"[PROMPT ACCESS] Admin check result: {is_admin}")
    
    if not is_admin:
        # Prompt is still loaded in background (module-level) based on role
        # but we don't return it to non-admin users
        logger.warning(f"[PROMPT ACCESS] Access denied for user: {username_header} (role: {user_role_header})")
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied. Admin role required to view prompt template."
        )
    
    # User is admin - return the role-specific prompt
    logger.info(f"[PROMPT ACCESS] Access granted for admin user: {username_header}")
    from .claude_client import _load_prompt_template_for_role
    # Get role from header to return appropriate prompt
    role = user_role_header if user_role_header != "NOT_SET" else ""
    prompt = _load_prompt_template_for_role(role)
    return {"prompt": prompt, "role": role}


@app.post("/api/detect-language")
async def detect_language(file: UploadFile = File(...)):
    """Detect the language of an uploaded document."""
    import tempfile
    
    try:
        # Save uploaded file temporarily
        file_content = await file.read()
        file_suffix = Path(file.filename).suffix if file.filename else ".pdf"
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=file_suffix)
        temp_file.write(file_content)
        temp_file.close()
        temp_path = Path(temp_file.name)
        
        try:
            detected_lang = detect_language_from_file(temp_path)
            return {"detected_language": detected_lang}
        finally:
            # Clean up temp file
            if temp_path.exists():
                temp_path.unlink()
    except Exception as e:
        logger.error(f"Language detection failed: {e}", exc_info=True)
        # Return default language on error
        return {"detected_language": "fr"}


@app.post("/api/translate", status_code=status.HTTP_202_ACCEPTED)
async def start_translation(
    request: Request,
    file: UploadFile = File(...),
    source_lang: str = Form("fr"),
    target_lang: str = Form("it"),
    use_glossary: bool = Form(False),
    skip_memory: bool = Form(False),  # Default False means memory is enabled
    custom_prompt: Optional[str] = Form(None),
    reference_doc: Optional[UploadFile] = File(None),
):
    """Start translation job. Each job is completely isolated."""
    import time
    
    # Get user role from headers for role-specific prompts
    user_role = (request.headers.get("X-User-Role") or request.headers.get("x-user-role") or "").strip()

    # Track this request
    request_timestamp = time.time()
    request_id = str(uuid.uuid4())[:8]
    
    logger.info("=" * 80)
    logger.info(f"[REQUEST {request_id}] ===== NEW TRANSLATION REQUEST =====")
    logger.info(f"[REQUEST {request_id}] Timestamp: {request_timestamp}")
    logger.info(f"[REQUEST {request_id}] File name: {file.filename}")
    logger.info(f"[REQUEST {request_id}] Source: {source_lang} -> Target: {target_lang}")
    
    settings = get_settings()
    if not settings.anthropic_api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured")
    
    # Read uploaded file content
    file_content = await file.read()
    file_hash = hashlib.md5(file_content).hexdigest()
    file_suffix = Path(file.filename).suffix if file.filename else ".pdf"
    
    logger.info(f"[REQUEST {request_id}] File size: {len(file_content):,} bytes")
    logger.info(f"[REQUEST {request_id}] File hash: {file_hash}")
    logger.info(f"[REQUEST {request_id}] First 200 bytes: {file_content[:200]}")
    logger.info(f"[REQUEST {request_id}] Use glossary: {use_glossary}")
    
    # Determine glossary path - use default if enabled
    glossary_path = None
    if use_glossary:
        # Use default glossary: glossary/glossary.csv
        default_glossary = Path("glossary") / "glossary.csv"
        if default_glossary.exists():
            glossary_path = str(default_glossary)
            logger.info(f"[REQUEST {request_id}] Using default glossary: {glossary_path}")
        else:
            logger.warning(f"[REQUEST {request_id}] Glossary enabled but default file not found: {default_glossary}")
    
    # Check for duplicate requests with same file hash - only reject if there's an ACTIVE job
    # Allow re-translation of completed jobs (they will use memory)
    recent_requests = [r for r in request_log if time.time() - r["timestamp"] < 60]  # Last 60 seconds
    duplicate_requests = [r for r in recent_requests if r["file_hash"] == file_hash]
    
    if duplicate_requests:
        # Check if any of the duplicate requests have active jobs
        active_duplicates = []
        for dup in duplicate_requests:
            dup_job_id = dup.get('job_id')
            if dup_job_id and dup_job_id in translation_jobs:
                job_status = translation_jobs[dup_job_id].get("status", "unknown")
                if job_status in ["pending", "in_progress"]:
                    active_duplicates.append((dup, job_status))
        
        if active_duplicates:
            logger.warning(f"[REQUEST {request_id}] ⚠️  DUPLICATE REQUEST DETECTED!")
            logger.warning(f"[REQUEST {request_id}] Found {len(active_duplicates)} active jobs with same file hash:")
            for dup, status in active_duplicates:
                logger.warning(f"[REQUEST {request_id}]   - Request {dup['request_id']} at {dup['timestamp']}, job_id: {dup.get('job_id')}, status: {status}")
            logger.warning(f"[REQUEST {request_id}] Active translation in progress - REJECTING duplicate")
            raise HTTPException(
                status_code=400,
                detail=f"Duplicate request detected. A translation with this file is already in progress."
            )
        else:
            # All duplicate requests are completed/failed - allow re-translation (will use memory)
            logger.info(f"[REQUEST {request_id}] Found {len(duplicate_requests)} recent requests with same hash, but all are completed")
            logger.info(f"[REQUEST {request_id}] Allowing re-translation - will check memory first")
    
    # Generate unique job ID
    job_id = str(uuid.uuid4())
    
    logger.info(f"[REQUEST {request_id}] Generated job_id: {job_id}")
    
    # Log this request
    request_log.append({
        "request_id": request_id,
        "job_id": job_id,
        "timestamp": request_timestamp,
        "file_hash": file_hash,
        "file_name": file.filename,
        "file_size": len(file_content),
    })
    
    # Keep only last 100 requests
    if len(request_log) > 100:
        request_log.pop(0)
    
    # Create unique temporary file for this job only
    # Use job_id in filename to ensure uniqueness
    temp_input = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=file_suffix,
        prefix=f"input_{job_id}_"
    )
    temp_input.write(file_content)
    temp_input.close()
    input_path = Path(temp_input.name)
    
    logger.info(f"Job {job_id}: Created isolated input file: {input_path.name}")
    logger.info(f"Job {job_id}: File hash: {file_hash[:16]}...")
    
    # Handle reference document if provided
    reference_doc_path = None
    if reference_doc:
        ref_doc_content = await reference_doc.read()
        ref_doc_suffix = Path(reference_doc.filename).suffix if reference_doc.filename else ".pdf"
        temp_ref = tempfile.NamedTemporaryFile(
            delete=False,
            suffix=ref_doc_suffix,
            prefix=f"reference_{job_id}_"
        )
        temp_ref.write(ref_doc_content)
        temp_ref.close()
        reference_doc_path = Path(temp_ref.name)
        logger.info(f"Job {job_id}: Reference document uploaded: {reference_doc_path.name}")
        logger.info(f"Job {job_id}: Reference doc size: {len(ref_doc_content):,} bytes")
    
    # Store ALL job data in job storage - no closure variables
    # CRITICAL: Check if job_id already exists (should never happen, but protect against it)
    if job_id in translation_jobs:
        logger.error(f"[REQUEST {request_id}] ⚠️  JOB ID COLLISION! Job {job_id} already exists!")
        logger.error(f"[REQUEST {request_id}] Existing job hash: {translation_jobs[job_id].get('file_hash')}")
        logger.error(f"[REQUEST {request_id}] New request hash: {file_hash}")
        raise HTTPException(status_code=500, detail="Job ID collision - please try again")
    
    translation_jobs[job_id] = {
        "status": "pending",
        "file_hash": file_hash,
        "input_path": str(input_path),
        "source_lang": source_lang,
        "target_lang": target_lang,
        "glossary_path": glossary_path,
        "use_glossary": use_glossary,
        "skip_memory": skip_memory,
        "custom_prompt": custom_prompt,
        "user_role": user_role,  # Store user role for role-specific prompts
        "reference_doc_path": str(reference_doc_path) if reference_doc_path else None,
        "progress": 0.0,
        "current_paragraph": 0,
        "total_paragraphs": 0,
        "cancelled": False,
        "request_id": request_id,  # Track which request created this job
        "created_at": request_timestamp,
    }
    
    logger.info(f"[REQUEST {request_id}] Job {job_id} stored in translation_jobs")
    logger.info(f"[REQUEST {request_id}] Total active jobs: {len([j for j in translation_jobs.values() if j.get('status') in ['pending', 'in_progress']])}")
    
    # Start translation thread - passes only job_id, reads everything from storage
    thread = threading.Thread(target=_run_translation, args=(job_id,), daemon=True)
    thread.start()
    
    logger.info(f"[REQUEST {request_id}] Job {job_id}: Translation thread started")
    logger.info(f"[REQUEST {request_id}] ===== REQUEST PROCESSING COMPLETE =====")
    logger.info("=" * 80)
    
    return {"job_id": job_id, "status": "pending"}


@app.get("/api/translate/{job_id}/status")
async def get_translation_status(job_id: str):
    if job_id not in translation_jobs:
        raise HTTPException(status_code=404, detail=JOB_NOT_FOUND)
    
    job = translation_jobs[job_id]
    return TranslationStatus(
        job_id=job_id,
        status=job["status"],
        progress=job.get("progress"),
        current_paragraph=job.get("current_paragraph"),
        total_paragraphs=job.get("total_paragraphs"),
        error=job.get("error"),
        pdf_size=job.get("pdf_size"),
        original_text=job.get("original_text"),
        translated_text=job.get("translated_text"),
        report=job.get("report"),
    )


@app.get("/api/translate/{job_id}/download")
async def download_translation(job_id: str):
    from fastapi.responses import Response
    
    if job_id not in translation_jobs:
        raise HTTPException(status_code=404, detail=JOB_NOT_FOUND)
    
    job = translation_jobs[job_id]
    if job["status"] != "completed":
        raise HTTPException(status_code=400, detail="Translation not completed")
    
    # Get PDF bytes from memory
    pdf_bytes = job.get("pdf_bytes")
    if not pdf_bytes:
        raise HTTPException(status_code=404, detail="PDF not found in memory")
    
    logger.info(f"[DOWNLOAD {job_id}] Serving PDF from memory: {len(pdf_bytes):,} bytes")
    
    # Generate filename
    source_lang = job.get("source_lang", "unknown")
    target_lang = job.get("target_lang", "unknown")
    filename = f"translated_{source_lang}_{target_lang}_{job_id[:8]}.pdf"
    
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Length": str(len(pdf_bytes)),
        }
    )


@app.post("/api/translate-stream")
async def translate_stream(
    request: Request,
    file: UploadFile = File(...),
    source_lang: str = Form(...),
    target_lang: str = Form(...),
    use_glossary: bool = Form(True),
    skip_memory: bool = Form(False),
    custom_prompt: Optional[str] = Form(None),
):
    """
    Stream translation results in real-time using Server-Sent Events (SSE).

    This endpoint provides immediate feedback as translation progresses,
    making the user experience feel much faster than waiting for completion.
    """
    import asyncio
    from queue import Queue

    settings = get_settings()
    if not settings.anthropic_api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured")

    # Read uploaded file
    file_content = await file.read()
    file_suffix = Path(file.filename).suffix if file.filename else ".txt"

    # Create temp file
    with tempfile.NamedTemporaryFile(mode='wb', suffix=file_suffix, delete=False) as tmp_file:
        tmp_file.write(file_content)
        temp_path = Path(tmp_file.name)

    try:
        # Read document text
        from .processing import _read_file_as_text
        document_text = _read_file_as_text(temp_path)

        # Load glossary
        glossary = None
        if use_glossary:
            default_glossary = Path("glossary") / "glossary.csv"
            if default_glossary.exists():
                glossary = _load_glossary(default_glossary, source_lang, target_lang)

        # Load memory
        memory_file = settings.data_root / "memory.json"
        memory = TranslationMemory(memory_file)

        # Get user role
        user_role = (request.headers.get("X-User-Role") or "").strip()

        # Determine prompt
        from .claude_client import _load_prompt_template_for_role
        prompt_to_use = custom_prompt or _load_prompt_template_for_role(user_role)

        # Create translator
        translator = ClaudeTranslator(
            api_key=settings.anthropic_api_key,
            dry_run=False,
            custom_prompt_template=prompt_to_use,
        )

        # Get glossary matches and memory hits
        glossary_matches = glossary.matches_in_text(document_text) if glossary else []
        memory_hits = [] if skip_memory else memory.similar(document_text, source_lang, target_lang, limit=10, threshold=80.0)

        # Check for exact memory match (100% similarity)
        exact_match = None
        if not skip_memory:
            exact_match = memory.get(document_text.strip(), source_lang, target_lang)

        async def event_generator():
            """Generate Server-Sent Events for streaming translation."""
            try:
                # Send start event with metadata
                start_data = {
                    'type': 'start',
                    'message': 'Translation started...',
                    'using_memory': exact_match is not None,
                    'glossary_matches': len(glossary_matches),
                    'memory_hits': len(memory_hits),
                }
                yield f"data: {json.dumps(start_data)}\n\n"

                # If exact memory match exists, return it immediately
                if exact_match:
                    full_translation = exact_match.translated_text
                    # Stream the cached translation character by character for visual effect
                    chunk_size = 50
                    for i in range(0, len(full_translation), chunk_size):
                        chunk = full_translation[i:i + chunk_size]
                        yield f"data: {json.dumps({'type': 'chunk', 'text': chunk})}\n\n"
                        await asyncio.sleep(0.01)  # Small delay for visual streaming effect
                else:
                    # Stream translation from Claude API
                    full_translation = ""
                    for chunk in translator.translate_document_streaming(
                        document_text=document_text,
                        source_lang=source_lang,
                        target_lang=target_lang,
                        glossary_matches=glossary_matches,
                        memory_hits=memory_hits,
                    ):
                        full_translation += chunk
                        # Send chunk event
                        yield f"data: {json.dumps({'type': 'chunk', 'text': chunk})}\n\n"
                        await asyncio.sleep(0)  # Allow other tasks to run

                    # Store in memory if not skipping
                    if not skip_memory:
                        memory.record(document_text, full_translation, source_lang, target_lang, allow_long_entries=True)

                # Send completion event with metadata
                completion_data = {
                    'type': 'done',
                    'message': 'Translation complete',
                    'full_text': full_translation,
                    'stats': {
                        'used_memory': exact_match is not None,
                        'glossary_matches': len(glossary_matches),
                        'memory_hits': len(memory_hits),
                        'source_lang': source_lang,
                        'target_lang': target_lang,
                    }
                }
                yield f"data: {json.dumps(completion_data)}\n\n"

            except Exception as e:
                logger.error(f"Streaming error: {e}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
            finally:
                # Cleanup temp file
                temp_path.unlink(missing_ok=True)

        return StreamingResponse(
            event_generator(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",  # Disable nginx buffering
            }
        )

    except Exception as e:
        # Cleanup on error
        temp_path.unlink(missing_ok=True)
        logger.error(f"Streaming setup error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/translate/{job_id}/cancel")
async def cancel_translation(job_id: str):
    """Cancel an ongoing translation job."""
    if job_id not in translation_jobs:
        raise HTTPException(status_code=404, detail=JOB_NOT_FOUND)
    
    job = translation_jobs[job_id]
    if job["status"] not in ["pending", "in_progress"]:
        raise HTTPException(status_code=400, detail="Translation cannot be cancelled")
    
    logger.info(f"Job {job_id}: Cancellation requested")
    job["cancelled"] = True
    
    return {"message": "Cancellation requested", "job_id": job_id}


@app.get("/api/translate/{job_id}/report")
async def get_translation_report(job_id: str):
    if job_id not in translation_jobs:
        raise HTTPException(status_code=404, detail=JOB_NOT_FOUND)
    
    job = translation_jobs[job_id]
    if job["status"] != "completed":
        raise HTTPException(status_code=400, detail="Translation not completed")
    
    return job.get("report", {})


@app.get("/api/translate/{job_id}/text")
async def get_translated_text(job_id: str):
    """Get translated text as plain text."""
    from fastapi.responses import Response
    
    if job_id not in translation_jobs:
        raise HTTPException(status_code=404, detail=JOB_NOT_FOUND)
    
    job = translation_jobs[job_id]
    if job["status"] != "completed":
        raise HTTPException(status_code=400, detail="Translation not completed")
    
    translated_paragraphs = job.get("translated_text")
    if not translated_paragraphs:
        raise HTTPException(status_code=404, detail="Translated text not found")
    
    # Join paragraphs with double newlines
    text_content = "\n\n".join(translated_paragraphs)
    
    return Response(
        content=text_content,
        media_type="text/plain; charset=utf-8",
        headers={
            "Content-Disposition": f'attachment; filename="translated_{job_id[:8]}.txt"',
        }
    )


@app.get("/api/translate/{job_id}/export/tmx")
async def export_translation_memory_tmx(job_id: str):
    """Export translation memory as TMX file for Trados."""
    import tempfile
    
    if job_id not in translation_jobs:
        raise HTTPException(status_code=404, detail=JOB_NOT_FOUND)
    
    job = translation_jobs[job_id]
    if job["status"] != "completed":
        raise HTTPException(status_code=400, detail="Translation not completed")
    
    source_lang = job.get("source_lang", "fr")
    target_lang = job.get("target_lang", "it")
    
    # Load translation memory
    settings = get_settings()
    memory_file = settings.data_root / "memory.json"
    memory = TranslationMemory(memory_file)
    
    # Create temporary TMX file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".tmx")
    temp_path = Path(temp_file.name)
    temp_file.close()
    
    try:
        # Export to TMX
        export_tmx(memory, temp_path, source_lang, target_lang)
        
        # Generate filename
        filename = f"translation_memory_{source_lang}_{target_lang}_{job_id[:8]}.tmx"
        
        return FileResponse(
            path=temp_path,
            filename=filename,
            media_type="application/xml",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
            }
        )
    except Exception as e:
        # Clean up temp file on error
        if temp_path.exists():
            temp_path.unlink()
        logger.error(f"Error exporting TMX: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error exporting TMX: {str(e)}")


@app.get("/api/translate/{job_id}/export/tbx")
async def export_glossary_tbx(job_id: str):
    """Export glossary as TBX file for Trados MultiTerm."""
    import tempfile
    
    if job_id not in translation_jobs:
        raise HTTPException(status_code=404, detail=JOB_NOT_FOUND)
    
    job = translation_jobs[job_id]
    if job["status"] != "completed":
        raise HTTPException(status_code=400, detail="Translation not completed")
    
    source_lang = job.get("source_lang", "fr")
    target_lang = job.get("target_lang", "it")
    glossary_path_str = job.get("glossary_path")
    
    if not glossary_path_str:
        raise HTTPException(status_code=400, detail="No glossary was used for this translation")
    
    # Load glossary
    glossary_path = Path(glossary_path_str)
    if not glossary_path.exists():
        raise HTTPException(status_code=404, detail="Glossary file not found")
    
    glossary = _load_glossary(glossary_path, source_lang, target_lang)
    if not glossary:
        raise HTTPException(status_code=500, detail="Failed to load glossary")
    
    # Create temporary TBX file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".tbx")
    temp_path = Path(temp_file.name)
    temp_file.close()
    
    try:
        # Export to TBX
        export_tbx(glossary, temp_path)
        
        # Generate filename
        filename = f"glossary_{source_lang}_{target_lang}_{job_id[:8]}.tbx"
        
        return FileResponse(
            path=temp_path,
            filename=filename,
            media_type="application/xml",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
            }
        )
    except Exception as e:
        # Clean up temp file on error
        if temp_path.exists():
            temp_path.unlink()
        logger.error(f"Error exporting TBX: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error exporting TBX: {str(e)}")
